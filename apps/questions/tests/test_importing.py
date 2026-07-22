import io

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document
from pypdf import PdfWriter

from apps.questions.importing import (
    ImportExtractionError,
    create_import_batch,
    extract_questions_from_csv,
    extract_questions_from_docx,
    extract_questions_from_pdf,
    extract_questions_from_plain_text,
    generate_question_title,
    normalize_question_text,
)
from apps.questions.models import Question, QuestionImportBatch, TechnicalQuestion

pytestmark = pytest.mark.django_db


def test_plain_text_extraction_supports_numbered_lists():
    text = """
    1. Explain Two Sum.
    2. Tell me about a conflict.
       Include what you learned.
    3. Find the off-by-one bug.
    """

    questions = extract_questions_from_plain_text(text)

    assert questions == [
        "Explain Two Sum.",
        "Tell me about a conflict.\nInclude what you learned.",
        "Find the off-by-one bug.",
    ]


def test_plain_text_extraction_supports_blank_line_blocks():
    text = "Explain hash maps.\n\nExplain sliding windows.\n\nDescribe a failure."

    questions = extract_questions_from_plain_text(text)

    assert questions == [
        "Explain hash maps.",
        "Explain sliding windows.",
        "Describe a failure.",
    ]


def test_csv_extraction_uses_named_question_column():
    file_obj = io.BytesIO(b"title,question\nOne,Explain heaps\nTwo,Explain graphs\n")

    questions = extract_questions_from_csv(file_obj)

    assert questions == ["Explain heaps", "Explain graphs"]


def test_docx_extraction_reads_paragraphs():
    document = Document()
    document.add_paragraph("1. Explain Two Sum")
    document.add_paragraph("2. Explain Binary Search")
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    questions = extract_questions_from_docx(output)

    assert questions == ["Explain Two Sum", "Explain Binary Search"]


def test_scanned_or_image_only_pdf_is_rejected():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)

    with pytest.raises(ImportExtractionError, match="Scanned PDFs"):
        extract_questions_from_pdf(output)


def test_generated_title_is_short_and_editable():
    title = generate_question_title(
        "Given an array of integers and a target, return the two matching indices."
    )

    assert title == "Given an array of integers and a target, return the…"


def test_normalization_removes_numbering_and_whitespace():
    assert normalize_question_text("  12)   Explain   Two Sum?  ") == "explain two sum?"


def test_create_batch_marks_existing_library_duplicate(user):
    TechnicalQuestion.objects.create(
        owner=user,
        title="Two Sum",
        prompt="Explain Two Sum",
    )

    batch = create_import_batch(
        owner=user,
        default_question_type=Question.Type.TECHNICAL,
        paste_text="1. Explain Two Sum\n2. Explain Binary Search",
    )

    items = list(batch.items.all())
    assert batch.status == QuestionImportBatch.Status.READY_FOR_REVIEW
    assert items[0].duplicate_reason == "Already exists in your question library."
    assert not items[0].is_included
    assert items[1].duplicate_reason == ""
    assert items[1].is_included


def test_invalid_text_file_creates_failed_batch(user):
    upload = SimpleUploadedFile("questions.txt", b"\xff\xfe\x00")

    batch = create_import_batch(
        owner=user,
        default_question_type=Question.Type.TECHNICAL,
        upload=upload,
    )

    assert batch.status == QuestionImportBatch.Status.FAILED
    assert "UTF-8" in batch.failure_message
